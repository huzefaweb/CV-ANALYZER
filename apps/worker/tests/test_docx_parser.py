"""DOCX parser adapter tests (AC#1, AC#2)."""

from __future__ import annotations

import io

import pytest
from docx import Document

from src.adapters.docx_parser import parse_docx
from src.domain.parse_gates import ParseFatalError


def _docx_bytes_from(build) -> bytes:
    doc = Document()
    build(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_text_native_paragraphs_extract_in_document_order():
    data = _docx_bytes_from(
        lambda doc: (
            doc.add_paragraph("Backend engineer for 5 years, primarily in Python and Django."),
            doc.add_paragraph("Built internal tools using PostgreSQL for caching."),
        )
    )
    units = parse_docx(data)
    assert [u.locator.path for u in units] == ["body/p[1]", "body/p[2]"]
    assert "Backend engineer" in units[0].text
    assert "PostgreSQL" in units[1].text


def test_table_cells_produce_stable_table_row_cell_paragraph_path():
    def build(doc):
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Skill"
        table.cell(0, 1).text = "Python"
        table.cell(1, 0).text = "Skill"
        table.cell(1, 1).text = "PostgreSQL"

    data = _docx_bytes_from(build)
    units = parse_docx(data)
    paths = [u.locator.path for u in units]
    assert "table[1]/row[1]/cell[1]/p[1]" in paths
    assert "table[1]/row[2]/cell[2]/p[1]" in paths


def test_body_paragraphs_and_tables_interleave_in_document_order():
    def build(doc):
        doc.add_paragraph("Employment history:")
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "Backend Engineer, Acme Corp"
        doc.add_paragraph("Education:")

    data = _docx_bytes_from(build)
    units = parse_docx(data)
    assert units[0].locator.path == "body/p[1]"
    assert units[1].locator.path == "table[1]/row[1]/cell[1]/p[1]"
    assert units[2].locator.path == "body/p[2]"


def test_nested_table_inside_a_cell_is_not_silently_dropped():
    def build(doc):
        table = doc.add_table(rows=1, cols=1)
        cell = table.cell(0, 0)
        cell.text = "Outer cell text"
        nested = cell.add_table(rows=1, cols=1)
        nested.cell(0, 0).text = "Nested cell text"

    data = _docx_bytes_from(build)
    units = parse_docx(data)
    texts = [u.text for u in units]
    assert "Outer cell text" in texts
    assert "Nested cell text" in texts
    nested_unit = next(u for u in units if u.text == "Nested cell text")
    assert "table[1]/row[1]/cell[1]/table[1]" in nested_unit.locator.path


def test_empty_paragraphs_are_skipped_not_fabricated():
    data = _docx_bytes_from(
        lambda doc: (doc.add_paragraph(""), doc.add_paragraph("Real content here."))
    )
    units = parse_docx(data)
    assert len(units) == 1
    assert units[0].text == "Real content here."


def test_corrupt_docx_raises_parse_fatal():
    with pytest.raises(ParseFatalError):
        parse_docx(b"this is not a docx file at all")


def test_empty_document_with_no_text_raises_parse_fatal():
    data = _docx_bytes_from(lambda doc: None)
    with pytest.raises(ParseFatalError):
        parse_docx(data)
